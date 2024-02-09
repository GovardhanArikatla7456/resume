from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from .models import Project
from django.template.loader import render_to_string
from django.core.files.base import ContentFile
import tempfile

# Create your views here.
def project_list(request):
    projects = Project.objects.all()  # Retrieve all projects
    return render(request, 'resumebuild/project_list.html', {'projects': projects})

def filter_projects(request):
    if request.method == 'POST':
        category = request.POST["category"]
        if category:
            projects = Project.objects.filter(category=category)
        else:
            projects = Project.objects.all()
        return render(request, 'resumebuild/project_list.html', {'projects': projects})
    else:
        return render(request, 'resumebuild/project_list.html')

def generate_resume(request):
    selected_projects = request.GET.getlist('selected_projects')
    projects = Project.objects.filter(id__in=selected_projects)
    print(selected_projects)
    # Create a new Docx document
    doc = Document()
    
    # Add the details of selected projects to the document
    for project in projects:
        doc.add_paragraph(f"Project Title: {project.name}")
        doc.add_paragraph(f"Description: {project.description}")
        doc.add_paragraph(f"Category: {project.category}")
        doc.add_paragraph("")  # Add an empty paragraph for separation
    
    # Create a temporary file to save the document
    temp_file = tempfile.NamedTemporaryFile(suffix=".docx")
    doc.save(temp_file.name)
    temp_file.seek(0)
    
    # Read the content of the temporary file
    file_content = temp_file.read()
    temp_file.close()
    
    # Generate the filename for the resume
    filename = 'resume.docx'
    
    # Save the Docx document to a response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    response.write(file_content)
    return response
